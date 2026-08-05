"""
Direct Labour Consult — CV Diagnostic Backend  V15B
====================================================
Render.com deployment  |  Python 3.11+
"""

print("=== DLC BACKEND V15B ACTIVE ===")

import gc
import os as _os

for _pv in ("HTTP_PROXY", "HTTPS_PROXY", "http_proxy", "https_proxy",
            "ALL_PROXY", "all_proxy", "NO_PROXY", "no_proxy"):
    _os.environ.pop(_pv, None)
import io
import json
import logging
import os
import re
import time
import unicodedata
from datetime import datetime, timezone
from pathlib import Path

from anthropic import Anthropic
from flask import Flask, jsonify, redirect, request
from flask_cors import CORS

try:
    import fitz
    PYMUPDF_OK = True
except ImportError:
    PYMUPDF_OK = False

try:
    from pdfminer.high_level import extract_text as pdfminer_extract
    PDFMINER_OK = True
except ImportError:
    PDFMINER_OK = False

try:
    import pytesseract
    from pdf2image import convert_from_bytes
    OCR_OK = True
except ImportError:
    OCR_OK = False

try:
    from docx import Document as DocxDocument
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("app")

app = Flask(__name__)

CORS(app, resources={r"/*": {
    "origins": [
        "https://directlabourconsult.com",
        "https://www.directlabourconsult.com",
        "https://dlc-cv-backend.onrender.com",
        "https://*.pages.dev",
        "http://localhost:3000",
        "http://localhost:5000",
        "http://127.0.0.1:5000",
    ],
    "expose_headers": ["X-DLC-Meta"],
}})

MAX_FILE_BYTES      = 10 * 1024 * 1024
MAX_OCR_PAGES       = 8
MIN_TEXT_CHARS      = 120
MAX_TEXT_CHARS      = 18_000
OCR_DPI             = 250
MAX_BATCH_FILES     = 25
MAX_TOKENS_SINGLE   = 4096
MAX_TOKENS_BATCH    = 3072
MAX_TOKENS_ATS      = 4096

ALLOWED_EXTENSIONS  = {"pdf", "doc", "docx"}
ANTHROPIC_MODEL     = "claude-opus-4-5"

BACKEND_VERSION     = "DLC_BACKEND_V15B"

ADMIN_KEY = os.environ.get("ADMIN_KEY", "dlc-admin-2026")

def check_admin(req) -> bool:
    key = req.headers.get("X-Admin-Key") or req.args.get("admin") or req.args.get("key") or req.form.get("admin_key")
    return bool(key and key == ADMIN_KEY)


def _build_anthropic_client() -> Anthropic:
    api_key = os.environ.get("ANTHROPIC_API_KEY", "").strip()
    if not api_key:
        raise EnvironmentError(
            "ANTHROPIC_API_KEY environment variable is not set. "
            "Add it in Render → Environment."
        )
    return Anthropic(api_key=api_key)


try:
    _client = _build_anthropic_client()
    log.info("Anthropic client initialised OK  model=%s", ANTHROPIC_MODEL)
except EnvironmentError as _e:
    _client = None
    log.error("Anthropic client NOT ready: %s", _e)


def _clean(raw: str) -> str:
    if not raw:
        return ""
    text = unicodedata.normalize("NFKD", raw)
    text = text.replace("\f", "\n").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    lines = [l for l in text.split("\n") if len(l.strip()) >= 2 or l.strip() == ""]
    return "\n".join(lines).strip()


def _sufficient(text: str) -> bool:
    return bool(text) and len(re.sub(r"\s+", "", text)) >= MIN_TEXT_CHARS


def _try_pymupdf(data: bytes) -> str:
    if not PYMUPDF_OK:
        return ""
    try:
        doc = fitz.open(stream=data, filetype="pdf")
        text = "\n".join(page.get_text("text") for page in doc)
        doc.close()
        return _clean(text)
    except Exception as exc:
        log.warning("PyMuPDF failed: %s", exc)
        return ""


def _try_pdfminer(data: bytes) -> str:
    if not PDFMINER_OK:
        return ""
    try:
        text = pdfminer_extract(io.BytesIO(data)) or ""
        return _clean(text)
    except Exception as exc:
        log.warning("pdfminer failed: %s", exc)
        return ""


def _try_ocr(data: bytes) -> tuple:
    if not OCR_OK:
        raise RuntimeError("OCR not available.")
    log.info("OCR starting …")
    try:
        images = convert_from_bytes(
            data, dpi=OCR_DPI, first_page=1, last_page=MAX_OCR_PAGES
        )
    except Exception as exc:
        raise RuntimeError(f"pdf2image conversion failed: {exc}") from exc

    parts = []
    done = 0
    for img in images:
        try:
            parts.append(pytesseract.image_to_string(img, lang="eng",
                                                      config="--oem 1 --psm 3"))
            done += 1
        except Exception as exc:
            log.warning("Tesseract error page %d: %s", done + 1, exc)
        finally:
            img.close()
            gc.collect()

    return _clean("\n".join(parts)), done


def _try_docx(data: bytes) -> str:
    if not DOCX_OK:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return _clean("\n".join(parts))
    except Exception as exc:
        log.warning("python-docx failed: %s", exc)
        return ""


def _try_doc_legacy(data: bytes) -> str:
    try:
        text = "".join(
            chr(b) if 32 <= b < 127 else (" " if b in (10, 13) else "")
            for b in data
        )
        runs = re.findall(r"[A-Za-z][A-Za-z0-9 \-\.,@:/()\n]{3,}", text)
        return _clean(" ".join(runs))
    except Exception as exc:
        log.warning("DOC legacy parser failed: %s", exc)
        return ""


def extract_cv_text(filename: str, data: bytes) -> tuple:
    ext = Path(filename).suffix.lower().lstrip(".")

    if ext == "pdf":
        text = _try_pymupdf(data)
        if _sufficient(text):
            log.info("Extraction: PyMuPDF  chars=%d", len(text))
            return text, "pymupdf"

        text = _try_pdfminer(data)
        if _sufficient(text):
            log.info("Extraction: pdfminer  chars=%d", len(text))
            return text, "pdfminer"

        log.info("Digital extraction insufficient (%d chars) — attempting OCR", len(text))
        try:
            text, pages = _try_ocr(data)
            if _sufficient(text):
                log.info("Extraction: OCR  pages=%d  chars=%d", pages, len(text))
                return text, "ocr ({} pages)".format(pages)
        except RuntimeError as exc:
            log.warning("OCR unavailable: %s", exc)

        raise ValueError(
            "We could not extract readable text from this PDF. "
            "Please save your CV directly from Microsoft Word or Google Docs "
            "(File → Save as PDF) and try again."
        )

    if ext == "docx":
        text = _try_docx(data)
        if _sufficient(text):
            log.info("Extraction: python-docx  chars=%d", len(text))
            return text, "python-docx"
        raise ValueError(
            "Could not read your DOCX file. "
            "Please re-save it in Microsoft Word and try again."
        )

    if ext == "doc":
        text = _try_doc_legacy(data)
        if _sufficient(text):
            log.info("Extraction: doc-legacy  chars=%d", len(text))
            return text, "doc-legacy"
        raise ValueError(
            "Could not read your DOC file. "
            "Please open it in Word and save as DOCX or PDF, then try again."
        )

    raise ValueError("Unsupported file type: .{}".format(ext))


def _strip_control_chars(s: str) -> str:
    return re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", s)


def _strip_markdown_fences(s: str) -> str:
    s = re.sub(r"^```(?:json)?\s*\n?", "", s.strip(), flags=re.IGNORECASE)
    s = re.sub(r"\n?```\s*$", "", s)
    return s.strip()


def _extract_json_object(s: str) -> str:
    start = s.find("{")
    if start == -1:
        return s

    depth = 0
    in_string = False
    escape_next = False

    for i, ch in enumerate(s[start:], start=start):
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s[start:i + 1]

    return s[start:]


def _repair_truncated_json(s: str) -> str:
    s = s.rstrip()
    s = re.sub(r",\s*$", "", s)

    depth_obj = 0
    depth_arr = 0
    in_string = False
    escape_next = False

    for ch in s:
        if escape_next:
            escape_next = False
            continue
        if ch == "\\" and in_string:
            escape_next = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == "{":
            depth_obj += 1
        elif ch == "}":
            depth_obj -= 1
        elif ch == "[":
            depth_arr += 1
        elif ch == "]":
            depth_arr -= 1

    if in_string:
        s += '"'

    s += "]" * max(0, depth_arr)
    s += "}" * max(0, depth_obj)

    return s


def _build_fallback_result(raw: str) -> dict:
    def _grab(pattern: str, default):
        m = re.search(pattern, raw, re.IGNORECASE | re.DOTALL)
        if m:
            try:
                return json.loads(m.group(1))
            except Exception:
                return m.group(1).strip().strip('"')
        return default

    score = _grab(r'"overall_score"\s*:\s*(\d+)', 0)
    name  = _grab(r'"candidate_name"\s*:\s*("(?:[^"\\]|\\.)*")', "Unknown")
    summary = _grab(r'"executive_summary"\s*:\s*("(?:[^"\\]|\\.)*")',
                    "Analysis completed. Full report unavailable — please retry.")

    return {
        "candidate_name":       name,
        "overall_score":        int(score) if str(score).isdigit() else 0,
        "market_readiness":     "Developing",
        "executive_summary":    summary,
        "top_strengths":        ["Analysis partially available — please retry for full report."],
        "critical_improvements":["Please retry for full diagnostic."],
        "sections": {
            "first_impression":   {"score": 0, "feedback": "Partial result — please retry.",
                                   "strengths": [], "improvements": []},
            "value_signal":       {"score": 0, "feedback": "Partial result — please retry.",
                                   "strengths": [], "improvements": []},
            "evidence_of_impact": {"score": 0, "feedback": "Partial result — please retry.",
                                   "strengths": [], "improvements": []},
            "role_alignment":     {"score": 0, "feedback": "Partial result — please retry.",
                                   "strengths": [], "improvements": []},
            "ats_compatibility":  {"score": 0, "feedback": "Partial result — please retry.",
                                   "strengths": [], "improvements": []},
        },
        "rewritten_section": {
            "section_name":     "Note",
            "original_excerpt": "",
            "rewritten":        "Full rewrite unavailable — please retry for complete report.",
        },
        "advisory_note": "Analysis was incomplete due to a response length issue. Please retry — your submission is valid.",
        "_partial": True,
    }


def safe_parse_json_v2(raw: str, context: str = "") -> dict:
    if not raw or not raw.strip():
        log.error("safe_parse_json_v2[%s]: empty response from Claude", context)
        return _build_fallback_result("")

    cleaned = _strip_markdown_fences(raw)
    cleaned = _strip_control_chars(cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    try:
        candidate = _extract_json_object(cleaned)
        return json.loads(candidate)
    except (json.JSONDecodeError, Exception):
        pass

    try:
        candidate = _extract_json_object(cleaned)
        repaired  = _repair_truncated_json(candidate)
        result    = json.loads(repaired)
        log.warning("safe_parse_json_v2[%s]: repaired truncated JSON", context)
        result["_repaired"] = True
        return result
    except (json.JSONDecodeError, Exception):
        pass

    log.error(
        "safe_parse_json_v2[%s]: all parse layers failed — using fallback. "
        "Raw first 400 chars: %s",
        context, cleaned[:400]
    )
    return _build_fallback_result(cleaned)


_PROMPT_SINGLE = """\
You are a senior HR advisory analyst at Direct Labour Consult (DLC), \
Botswana's premier HR consultancy (established 2018). \
Evaluate the CV below using DLC's five-layer diagnostic framework. \
Return ONLY a raw JSON object — no markdown, no code fences, no explanation text.

Required JSON structure (return ALL fields, be concise to stay within token limit):
{{
  "candidate_name":       "string — inferred from CV",
  "overall_score":        integer 0-100,
  "market_readiness":     "Excellent | Strong | Developing | Needs Improvement",
  "executive_summary":    "2-3 sentences. Specific, professional, honest.",
  "top_strengths":        ["string","string","string"],
  "critical_improvements":["string","string","string"],
  "sections": {{
    "first_impression":   {{"score":int,"feedback":"string max 80 words","strengths":["str"],"improvements":["str"]}},
    "value_signal":       {{"score":int,"feedback":"string max 80 words","strengths":["str"],"improvements":["str"]}},
    "evidence_of_impact": {{"score":int,"feedback":"string max 80 words","strengths":["str"],"improvements":["str"]}},
    "role_alignment":     {{"score":int,"feedback":"string max 80 words","strengths":["str"],"improvements":["str"]}},
    "ats_compatibility":  {{"score":int,"feedback":"string max 80 words","strengths":["str"],"improvements":["str"]}}
  }},
  "rewritten_section": {{
    "section_name":     "e.g. Professional Summary",
    "original_excerpt": "verbatim excerpt 40 words max",
    "rewritten":        "improved version 60 words max"
  }},
  "advisory_note": "1-2 sentences. Single most important action."
}}

Scoring: 85-100 Excellent | 70-84 Strong | 50-69 Developing | 0-49 Needs Improvement
Keep all feedback concise. Every observation must be traceable to actual CV content.

CV CONTENT:
---
{cv_text}
---"""


_PROMPT_BATCH_CANDIDATE = """\
You are a Senior Assessment Consultant at Direct Labour Consult (DLC), \
Botswana's premier Industrial Psychology-Informed Executive Assessment practice. \
Conduct a rigorous, evidence-based candidate evaluation for the recruitment mandate below. \
Your assessment must reflect the depth and analytical standard of a Senior Industrial/Organisational \
Psychology practitioner — not a generic screener. \
Return ONLY a valid JSON object — no markdown, no code fences, no preamble.

RECRUITMENT MANDATE:
  Position:        {job_title}
  Department:      {job_dept}
  Min. Experience: {job_exp} years
  Education:       {job_edu}
  Key Skills:      {job_skills}
  Job Description: {job_desc}
  Disqualifiers:   {job_disq}

ASSESSMENT FRAMEWORK — apply all five dimensions:
  1. COMPETENCY EVIDENCE     Distinguish demonstrated competencies from claimed ones. \
Cite specific employers, roles, tenures, and achievements from the CV.
  2. CAREER TRAJECTORY       Analyse progression patterns: ascending, lateral, or fragmented. \
Note scope of accountability growth and any tenure concerns.
  3. PERSON-JOB FIT          Calibrate directly against the role demands above. \
Flag alignment factors and material gaps specific to this mandate.
  4. BEHAVIOURAL INDICATORS  Infer motivational orientation, resilience signals, and \
cultural fit from CV structure, tone, content patterns, and career choices.
  5. OCCUPATIONAL MATURITY   Assess functional readiness relative to the seniority level required \
for this position.

REQUIRED JSON (all fields mandatory):
{{
  "candidate_name":      "Full name inferred from CV. Use filename stem if not found.",
  "overall_score":       integer 0-100,
  "recommendation":      "Hire | Consider | Not Aligned",

  "executive_summary":   "3 sentences. Open with the single most significant finding about this candidate's suitability. Reference specific CV evidence — employer names, role titles, tenures, or quantified achievements. Close with a decisive and consultative assessment stance. Authoritative tone — no generic filler.",

  "job_fit_note":        "1-2 sentences. Specific alignment or gap between this candidate's demonstrated experience and the role requirements stated above. Name the role and evidence.",

  "key_strengths": [
    "Evidence-anchored: cite specific employer / role / achievement from CV, then state what this indicates about capability for this role",
    "Evidence-anchored: cite specific evidence — then state what this indicates",
    "Evidence-anchored: cite specific evidence — then state what this indicates"
  ],
  "key_concerns": [
    "Name the specific gap or risk — cite absence of evidence or a problematic pattern observed in the CV",
    "Name the specific gap or risk — cite evidence"
  ],

  "behavioural_risk":   "Low | Medium | High",
  "behavioural_notes":  "Analytical observation on behavioural or motivational risk indicators. Reference tenure patterns, unexplained gaps, career move motivations, or stated interests as signals. Empty string if risk is Low.",

  "sections": {{
    "first_impression":   {{"score": integer 0-100, "rationale": "What does CV structure, formatting, and professional framing signal about this candidate's professional identity and self-presentation standard?"}},
    "evidence_of_impact": {{"score": integer 0-100, "rationale": "Quality and specificity of achievement evidence — are outcomes quantified, attributed, and credible? Or are responsibilities listed without impact?"}},
    "role_alignment":     {{"score": integer 0-100, "rationale": "Degree of match between demonstrated experience and the specific functional demands of this role."}},
    "ats_compatibility":  {{"score": integer 0-100, "rationale": "Keyword density, role-relevant terminology, and structural scannability for this type of position."}}
  }},

  "occupational_profile": {{
    "leadership_readiness":  "Emerging | Developing | Established | Advanced",
    "operational_maturity":  "Graduate | Junior | Mid-level | Senior | Executive",
    "strategic_thinking":    "Absent | Limited | Present | Strong",
    "stakeholder_exposure":  "Internal only | Cross-functional | External / Board-level"
  }},

  "competency_assessment": {{
    "technical_competence":      "Below | Meets | Exceeds — 1-sentence evidence statement citing specific CV content",
    "analytical_reasoning":      "Below | Meets | Exceeds — 1-sentence evidence statement",
    "communication_proficiency": "Below | Meets | Exceeds — 1-sentence evidence statement",
    "leadership_and_influence":  "Below | Meets | Exceeds — 1-sentence evidence statement"
  }},

  "career_trajectory":   "1-2 sentences. Is the career pattern ascending, lateral, or fragmented? What does the scope of accountability growth — or absence of it — indicate about this candidate's ceiling and professional drive?",

  "recruiter_guidance":  "Recommended for Interview | Recommended for Shortlist | Development Candidate | Not Aligned Currently",
  "years_experience":    "estimated range e.g. 5-7 years",
  "education_alignment": "Exceeds requirements | Meets requirements | Below requirements | Unable to determine",
  "advisory_note":       "1-2 sentences. The single most important nuanced observation for the hiring manager — something that standard CV screening would miss. May relate to a hidden strength, a structural risk, or a context-specific consideration."
}}

SCORING FRAMEWORK:
  80-100  Strong, well-evidenced profile — recommend Hire
  60-79   Viable candidate with reservations — recommend Consider
  0-59    Significant gaps against mandate requirements — recommend Not Aligned

QUALITY STANDARDS:
- Every observation must cite specific CV content: employer name, role title, tenure, metric, or named achievement
- No generic filler (avoid: "demonstrates strong experience", "has a proven track record", "brings valuable skills")
- Tone: authoritative, consultative, professionally direct — not robotic, not harsh
- Clearly distinguish demonstrated evidence from reasonable inferences
- If job description is blank, calibrate against established professional standards for this role type
- If candidate name is unclear, derive it from the filename

CV CONTENT:
---
{cv_text}
---"""


def analyse_with_claude(cv_text: str, name_hint: str = "") -> dict:
    if _client is None:
        raise RuntimeError(
            "Analysis service is not configured (missing API key). "
            "Contact DLC support."
        )

    if len(cv_text) > MAX_TEXT_CHARS:
        log.info("CV text truncated %d → %d chars", len(cv_text), MAX_TEXT_CHARS)
        cv_text = cv_text[:MAX_TEXT_CHARS] + "\n\n[Content truncated for analysis]"

    log.info("Single CV: %d chars", len(cv_text))

    msg = _client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=MAX_TOKENS_SINGLE,
        messages=[{
            "role":    "user",
            "content": _PROMPT_SINGLE.format(cv_text=cv_text),
        }],
    )

    raw = msg.content[0].text.strip()
    log.info("Claude response length: %d chars, stop_reason: %s",
             len(raw), msg.stop_reason)

    if msg.stop_reason == "max_tokens":
        log.warning("Claude hit max_tokens limit — response may be truncated")

    result = safe_parse_json_v2(raw, context="single_cv")

    if not result.get("candidate_name") and name_hint:
        result["candidate_name"] = name_hint

    return result


def analyse_batch_candidate(cv_text: str, filename: str, job_params: dict = None) -> dict:
    if _client is None:
        raise RuntimeError("Analysis service not configured.")

    if len(cv_text) > 10_000:
        cv_text = cv_text[:10_000] + "\n\n[Truncated]"

    p = job_params or {}
    prompt = _PROMPT_BATCH_CANDIDATE.format(
        cv_text   = cv_text,
        job_title = p.get("job_title", "Not specified"),
        job_dept  = p.get("job_dept",  "Not specified"),
        job_exp   = p.get("job_exp",   "Not specified"),
        job_edu   = p.get("job_edu",   "Not specified"),
        job_skills= p.get("job_skills","Not specified"),
        job_desc  = p.get("job_desc",  ""),
        job_disq  = p.get("job_disq",  "None stated"),
    )

    msg = _client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=MAX_TOKENS_BATCH,
        messages=[{
            "role":    "user",
            "content": prompt,
        }],
    )

    raw = msg.content[0].text.strip()
    result = safe_parse_json_v2(raw, context="batch_{}".format(filename))

    if not result.get("candidate_name"):
        result["candidate_name"] = (
            Path(filename).stem.replace("_", " ").replace("-", " ").title()
        )

    rec = (result.get("recommendation") or "").strip()
    if rec.lower() in ("hire", "strong hire"):
        result["recommendation"] = "Hire"
    elif rec.lower() in ("consider",):
        result["recommendation"] = "Consider"
    else:
        result["recommendation"] = "Not Aligned"

    br = (result.get("behavioural_risk") or "Low").strip()
    if "high" in br.lower():
        result["behavioural_risk"] = "High"
    elif "med" in br.lower():
        result["behavioural_risk"] = "Medium"
    else:
        result["behavioural_risk"] = "Low"

    sects = result.get("sections") or {}
    for sk in ("first_impression", "evidence_of_impact", "role_alignment", "ats_compatibility"):
        if sk not in sects or not isinstance(sects.get(sk), dict):
            sects[sk] = {"score": 0}
        elif "score" not in sects[sk]:
            sects[sk]["score"] = 0
    result["sections"] = sects

    if not result.get("key_strengths"):
        result["key_strengths"] = result.pop("top_strengths", []) or []
    if not result.get("key_concerns"):
        result["key_concerns"] = result.pop("critical_gaps", []) or []
    if not result.get("behavioural_notes"):
        result["behavioural_notes"] = result.pop("behavioural_risk_note", "") or ""
    if not result.get("executive_summary"):
        result["executive_summary"] = result.pop("recommendation_reason", "") or ""

    result["_filename"] = filename

    return result


def _get_batch_files(req):
    log.info("FILES RECEIVED: %s", list(req.files.keys()))
    log.info("FORM FIELDS: %s",    list(req.form.keys()))

    for field in ("files", "files[]", "cvs", "cvs[]", "cv_files", "cv_files[]"):
        batch = req.files.getlist(field)
        if batch and any(f.filename for f in batch):
            log.info("Batch files found under field=%r  count=%d", field, len(batch))
            return [f for f in batch if f.filename]

    all_files = []
    for key in req.files.keys():
        all_files.extend(req.files.getlist(key))
    all_files = [f for f in all_files if f.filename]

    if all_files:
        log.info("Batch files found via fallback scan  count=%d", len(all_files))

    return all_files


@app.route("/")
def index():
    return redirect("https://directlabourconsult.com", code=302)


@app.route("/health", methods=["GET"])
def health():
    return jsonify({
        "success": True,
        "status":  "ok",
        "version": BACKEND_VERSION,
        "model":   ANTHROPIC_MODEL,
        "capabilities": {
            "pymupdf":  PYMUPDF_OK,
            "pdfminer": PDFMINER_OK,
            "ocr":      OCR_OK,
            "docx":     DOCX_OK,
        },
    }), 200


@app.route("/version", methods=["GET"])
def version():
    return jsonify({
        "version":          BACKEND_VERSION,
        "parser":           "safe_parse_json_v2",
        "upload_detection": "multi-field",
        "max_tokens_single": MAX_TOKENS_SINGLE,
        "max_tokens_batch":  MAX_TOKENS_BATCH,
        "status":           "active",
        "model":            ANTHROPIC_MODEL,
    }), 200


@app.route("/analyze", methods=["POST", "OPTIONS"])
def analyze():
    if request.method == "OPTIONS":
        return "", 204

    t0 = time.time()
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    def err(msg: str, status: int = 400):
        log.warning("ANALYZE ERROR [%s]: %s", ts, msg)
        return jsonify({"success": False, "error": msg}), status

    name  = (request.form.get("name")  or "").strip()
    email = (request.form.get("email") or "").strip()

    if not name:
        return err("Name is required.")
    if not email or "@" not in email:
        return err("A valid email address is required.")

    cv_file = request.files.get("cv") or request.files.get("cv_original")
    if not cv_file or not cv_file.filename:
        return err("No CV file received. Please select and upload your CV.")

    filename = cv_file.filename
    ext      = Path(filename).suffix.lower().lstrip(".")

    if ext not in ALLOWED_EXTENSIONS:
        return err(
            "File type '.{}' is not supported. "
            "Please upload a PDF, DOCX, or DOC file.".format(ext)
        )

    try:
        data = cv_file.read()
    except Exception as exc:
        log.error("File read error: %s", exc)
        return err("Could not read the uploaded file. Please try again.", 500)

    if len(data) > MAX_FILE_BYTES:
        return err(
            "File size ({} MB) exceeds the 10 MB limit.".format(len(data) // 1024 // 1024)
        )
    if len(data) < 500:
        return err("The uploaded file appears to be empty or corrupt.")

    log.info("ANALYZE START  name=%r  file=%r  size=%d bytes", name, filename, len(data))

    method  = "unknown"
    cv_text = ""

    client_text = (request.form.get("extracted_text") or "").strip()
    if client_text and _sufficient(client_text):
        cv_text = _clean(client_text)
        method  = "client-preextracted"
        log.info("Using client pre-extracted text  chars=%d", len(cv_text))
    else:
        try:
            cv_text, method = extract_cv_text(filename, data)
        except ValueError as exc:
            log.warning("Extraction failed: %s", exc)
            return err(str(exc), 422)
        except Exception as exc:
            log.error("Unexpected extraction error: %s", exc, exc_info=True)
            return err(
                "An unexpected error occurred while reading your CV. "
                "Please try again or contact DLC support.", 500
            )

    try:
        result = analyse_with_claude(cv_text, name_hint=name)
    except RuntimeError as exc:
        log.error("Claude analysis error: %s", exc)
        return err(str(exc), 503)
    except Exception as exc:
        log.error("Unexpected Claude error: %s", exc, exc_info=True)
        return err(
            "Analysis service encountered an unexpected error. "
            "Your payment is still valid — please try again in a moment.", 503
        )

    elapsed = round(time.time() - t0, 2)
    result["_meta"] = {
        "extraction_method": method,
        "chars_analysed":    len(cv_text),
        "ocr_used":          "ocr" in method,
        "processing_time_s": elapsed,
        "timestamp_utc":     ts,
        "backend_version":   BACKEND_VERSION,
    }

    log.info(
        "ANALYZE DONE  name=%r  score=%s  method=%s  time=%.2fs",
        name, result.get("overall_score", "?"), method, elapsed,
    )

    return jsonify({"success": True, "data": result}), 200


@app.route("/analyze-batch", methods=["POST", "OPTIONS"])
def analyze_batch():
    if request.method == "OPTIONS":
        return "", 204

    t0 = time.time()
    ts = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    def err(msg: str, status: int = 400):
        log.warning("BATCH ERROR [%s]: %s", ts, msg)
        return jsonify({"success": False, "error": msg}), status

    job_params = {
        "job_title":  (request.form.get("job_title")  or "").strip(),
        "job_dept":   (request.form.get("job_dept")   or "").strip(),
        "job_exp":    (request.form.get("job_exp")    or "").strip(),
        "job_edu":    (request.form.get("job_edu")    or "").strip(),
        "job_skills": (request.form.get("job_skills") or "").strip(),
        "job_desc":   (request.form.get("job_desc")   or "").strip(),
        "job_disq":   (request.form.get("job_disq")   or "").strip(),
    }
    log.info(
        "BATCH JOB PARAMS  title=%r  dept=%r  exp=%r  edu=%r",
        job_params["job_title"], job_params["job_dept"],
        job_params["job_exp"],   job_params["job_edu"],
    )

    uploaded = _get_batch_files(request)

    if not uploaded:
        log.error(
            "BATCH: No files detected. "
            "Files dict keys: %s | Content-Type: %s",
            list(request.files.keys()),
            request.content_type,
        )
        return err(
            "No CV files received. Please ensure files are attached before submitting. "
            "Accepted formats: PDF, DOCX, DOC."
        )

    if len(uploaded) > MAX_BATCH_FILES:
        return err(
            "Maximum {} CVs per batch. You submitted {}.".format(
                MAX_BATCH_FILES, len(uploaded)
            )
        )

    valid_files = []
    skipped     = []
    for f in uploaded:
        ext = Path(f.filename).suffix.lower().lstrip(".")
        if ext in ALLOWED_EXTENSIONS:
            valid_files.append(f)
        else:
            skipped.append(f.filename)

    if skipped:
        log.warning("Batch: skipping unsupported files: %s", skipped)

    if not valid_files:
        return err(
            "None of the uploaded files are in a supported format. "
            "Please upload PDF, DOCX, or DOC files."
        )

    log.info("BATCH START  files=%d  valid=%d  ts=%s", len(uploaded), len(valid_files), ts)

    results = []
    errors  = []

    for i, cv_file in enumerate(valid_files):
        filename = cv_file.filename
        log.info("Batch processing [%d/%d]: %s", i + 1, len(valid_files), filename)

        try:
            data = cv_file.read()
        except Exception as exc:
            log.error("Batch file read error [%s]: %s", filename, exc)
            errors.append({"filename": filename, "error": "Could not read file."})
            continue

        if len(data) > MAX_FILE_BYTES:
            errors.append({"filename": filename, "error": "File too large (max 10 MB)."})
            continue

        if len(data) < 200:
            errors.append({"filename": filename, "error": "File appears empty."})
            continue

        try:
            cv_text, method = extract_cv_text(filename, data)
        except ValueError as exc:
            log.warning("Batch extraction failed [%s]: %s", filename, exc)
            errors.append({"filename": filename, "error": str(exc)})
            continue
        except Exception as exc:
            log.error("Batch extraction error [%s]: %s", filename, exc)
            errors.append({"filename": filename, "error": "Text extraction failed."})
            continue

        try:
            analysis = analyse_batch_candidate(cv_text, filename, job_params)
        except RuntimeError as exc:
            log.error("Batch Claude error [%s]: %s", filename, exc)
            errors.append({"filename": filename, "error": str(exc)})
            continue
        except Exception as exc:
            log.error("Batch unexpected Claude error [%s]: %s", filename, exc)
            errors.append({"filename": filename, "error": "Analysis failed."})
            continue

        analysis["_filename"]        = filename
        analysis["extraction_method"] = method
        results.append(analysis)

    if not results:
        return jsonify({
            "success": False,
            "error":   "Analysis failed for all submitted CVs.",
            "errors":  errors,
        }), 500

    results.sort(key=lambda x: x.get("overall_score", 0), reverse=True)
    for rank, r in enumerate(results, start=1):
        r["rank"] = rank

    hire_count    = sum(1 for r in results if r.get("recommendation") == "Hire")
    consider_count = sum(1 for r in results if r.get("recommendation") == "Consider")
    not_aligned   = sum(1 for r in results if r.get("recommendation") == "Not Aligned")
    scores        = [r.get("overall_score", 0) for r in results if isinstance(r.get("overall_score"), (int, float))]
    avg_score     = round(sum(scores) / len(scores)) if scores else 0
    top_score     = max(scores) if scores else 0

    if avg_score >= 72:
        pool_quality = "Strong candidate pool"
    elif avg_score >= 55:
        pool_quality = "Mixed candidate pool"
    else:
        pool_quality = "Thin candidate pool — consider widening the search"

    summary = {
        "hire":          hire_count,
        "consider":      consider_count,
        "not_aligned":   not_aligned,
        "average_score": avg_score,
        "top_score":     top_score,
        "pool_quality":  pool_quality,
    }

    top = results[0] if results else {}
    executive_report = (
        "{} candidate{} assessed. {} recommended for hire, {} for consideration. "
        "Average pool score: {}/100. Top candidate: {} ({}). {}".format(
            len(results),
            "s" if len(results) != 1 else "",
            hire_count,
            consider_count,
            avg_score,
            top.get("candidate_name", "N/A"),
            top.get("overall_score", "N/A"),
            pool_quality + ".",
        )
    )

    elapsed = round(time.time() - t0, 2)

    log.info(
        "BATCH DONE  processed=%d  hire=%d  consider=%d  avg=%d  time=%.2fs",
        len(results), hire_count, consider_count, avg_score, elapsed,
    )

    return jsonify({
        "success":           True,
        "ranked_candidates": results,
        "total_candidates":  len(results),
        "summary":           summary,
        "executive_report":  executive_report,
        "job_title":         job_params.get("job_title", ""),
        "total_submitted":   len(valid_files),
        "errors":            errors,
        "processing_time_s": elapsed,
        "timestamp_utc":     ts,
        "backend_version":   BACKEND_VERSION,
    }), 200


@app.errorhandler(404)
def not_found(_e):
    return jsonify({"success": False, "error": "Endpoint not found."}), 404


@app.errorhandler(405)
def method_not_allowed(_e):
    return jsonify({"success": False, "error": "Method not allowed."}), 405


@app.errorhandler(413)
def request_entity_too_large(_e):
    return jsonify({"success": False, "error": "File too large. Maximum size is 10 MB."}), 413


@app.errorhandler(500)
def internal_error(exc):
    log.error("Unhandled 500: %s", exc, exc_info=True)
    return jsonify({
        "success": False,
        "error": "An internal server error occurred. Please try again.",
    }), 500


# ══════════════════════════════════════════════════════════════════════════
# ── DLC ATS APPLICATION ENGINE — additive feature.
#    Ported onto the real V15B base above. Uses V15B's actual primitives
#    (extract_cv_text, _client, safe_parse_json_v2, log) rather than a
#    separate lineage. Nothing above this line was modified.
# ══════════════════════════════════════════════════════════════════════════

from flask import send_file
import base64
import smtplib
from email.message import EmailMessage

ATS_OPTIMIZE_SYSTEM = """You are the DLC ATS Application Engine™ — acting simultaneously as an Applicant Tracking System and a Chartered HR Practitioner / Occupational Psychologist at Direct Labour Consult. You evaluate one candidate against one specific job description, the way a real recruiter and a real ATS would, and then produce an optimised application.

Work through these steps using the CV (and, if provided, the target job description):

1. PARSE THE CV — extract the candidate's real name, most recent/likely professional title, location if stated, contact details if present, skills, full work history, education, and certifications. Never invent a name, employer, date, or qualification the candidate did not provide — use "[Not provided]" as a placeholder for header fields you cannot find, rather than guessing.
2. IF A JOB DESCRIPTION IS PROVIDED — identify required skills, preferred/nice-to-have skills, keywords, and the seniority/experience level implied, then align the rewritten CV to it.
3. MATCH (only if a job description was provided) — compare CV to role: skills overlap, experience alignment, keyword overlap, title alignment.
4. IDENTIFY GAPS (only if a job description was provided) — required skills missing or under-evidenced, weakly-written bullets, keywords the role uses that the CV doesn't.
5. IDENTIFY RISKS — flag realistic hiring-manager concerns ONLY when the evidence genuinely supports them (job hopping, unexplained gaps, over/underqualification, ATS-unfriendly formatting). Do not invent risks.
6. SCORE — produce these metrics (0-100 each unless noted), each grounded in real evidence from the CV, never inflated for effect:
   - ats_compatibility_score: how well-structured and machine-parseable the CV itself is.
   - job_match_score (or null): only if a job description was provided — how well the candidate matches THIS role.
   - formatting_quality_score: structural clarity — consistent sections, clean chronology, no ATS-breaking elements.
   - keyword_optimization_score (or null if no job description): how well role-relevant keywords are naturally represented.
   - professional_positioning_score: strength and clarity of the professional summary and title framing.
   - recruiter_readiness_score: how quickly a recruiter could assess fit within a 7-second scan.
   - vacancy_alignment_score (or null if no job description): overall fit against the specific vacancy's stated requirements.
   - evidence_strength_score: how well claims are backed by specifics (employers, tenures, quantified outcomes) rather than vague responsibility statements.
   - hiring_readiness_score: weighted composite headline score. If a job description was provided: ATS Compatibility 25% / Job Match 40% / Content Strength 20% / Risk Factors 15%. If no job description was provided, reweight as: ATS Compatibility 40% / Content Strength 40% / Risk Factors 20%.
7. BUILD THE FULL CV — produce a complete, corporate, ATS-compliant CV using ONLY the candidate's real, provided experience and qualifications:
   - Professional summary: 4-5 lines, strong corporate language, naming years of experience (estimate conservatively from the CV's own dates if not explicit), industry focus, key strengths, and value to an employer. Do NOT use generic filler phrases such as "results-driven professional," "team player," "dynamic professional," "passionate about," or "synergy."
   - Core competencies: 8-12 concise ATS keyword phrases genuinely evidenced by the candidate's real background — never invent expertise the CV doesn't support.
   - Professional experience: for EACH role found in the CV, rewrite 4-6 achievement-based bullets using strong action verbs (Led, Implemented, Managed, Improved, Delivered, Negotiated, Streamlined). Never fabricate metrics, outcomes, or responsibilities not implied by the original content.
   - Education: cleanly rewritten from what the CV states.
   - Certifications: list only what the CV or supporting documents actually evidence. Return an empty list if none are found.
   - References: if the candidate's CV lists actual named referees (name, title/relationship, organisation, and/or contact details), include exactly what is stated — never invent names, titles, organisations, or contact details. If the CV does not list named referees, set "references_available_on_request" to true and leave the references list empty — do NOT fabricate placeholder people.
8. RECRUITER'S FIRST IMPRESSION — write 2-3 sentences in the voice of a senior HR consultant giving a genuine, candid first read of this candidate, as if briefing a hiring manager. Name the specific professional identity this CV projects, note one genuine strength, and end with one concrete, specific suggestion for what would strengthen future applications (e.g. a category of evidence that's currently thin — quantified outcomes, budget/portfolio size, named methodologies). This must read as genuine consultative judgment, not a restatement of the summary, and must be grounded only in what the CV actually shows.

Use professional HR language throughout.

PROHIBITED WORDS: suffers, corrupted, failure, rejected, unreadable, terrible, poor, bad, wrong, weak, broken, results-driven, synergy.

Return ONLY a valid JSON object — no markdown, no preamble, no commentary:

{
  "ats_score": <0-100 integer>,
  "job_match_score": <0-100 integer, or null if no job description was provided>,
  "formatting_quality_score": <0-100 integer>,
  "keyword_optimization_score": <0-100 integer, or null if no job description was provided>,
  "professional_positioning_score": <0-100 integer>,
  "recruiter_readiness_score": <0-100 integer>,
  "vacancy_alignment_score": <0-100 integer, or null if no job description was provided>,
  "evidence_strength_score": <0-100 integer>,
  "hiring_readiness_score": <0-100 integer>,
  "gaps": ["<gap 1>", "<gap 2>"],
  "risks": ["<risk, only if genuinely evidenced — empty array if none>"],
  "recruiters_first_impression": "<2-3 sentence senior-consultant narrative, per step 8>",
  "cv_document": {
    "header": {
      "full_name": "<candidate's real name, or '[Not provided]'>",
      "professional_title": "<derived from most recent/dominant experience>",
      "location": "<if stated, else '[Not provided]'>",
      "email": "<if stated, else '[Not provided]'>",
      "phone": "<if stated, else '[Not provided]'>"
    },
    "professional_summary": "<4-5 line corporate summary>",
    "core_competencies": ["<competency 1>", "... 8-12 total"],
    "professional_experience": [
      {
        "title": "<role title>",
        "employer": "<employer name>",
        "dates": "<date range as stated>",
        "bullets": ["<achievement bullet 1>", "... 4-6 total"]
      }
    ],
    "education": [
      {"qualification": "<degree/qualification>", "institution": "<institution>", "dates": "<if stated>"}
    ],
    "certifications": ["<certification 1>", "... or empty array if none found"],
    "references": [
      {"name": "<referee's real name, only if stated>", "title": "<referee's title/relationship, if stated>", "organisation": "<if stated>", "contact": "<if stated>"}
    ],
    "references_available_on_request": <true if no named referees were found in the CV, otherwise false>
  },
  "keywords_added": ["<keyword 1>", "<keyword 2>"]
}"""


def call_claude_ats(system_prompt: str, user_content: str, image_blocks=None) -> dict:
    """
    ATS-specific Claude call helper. V15B has no existing call_claude wrapper,
    so this is new — but reuses V15B's real _client, ANTHROPIC_MODEL, and
    safe_parse_json_v2 (which never raises) rather than introducing a
    second, incompatible JSON parser.
    """
    if _client is None:
        raise RuntimeError("Analysis service is not configured (missing API key).")

    content = ([{"type": "text", "text": user_content}] + image_blocks) if image_blocks else user_content
    msg = _client.messages.create(
        model=ANTHROPIC_MODEL,
        max_tokens=MAX_TOKENS_ATS,
        system=system_prompt,
        messages=[{"role": "user", "content": content}],
    )
    raw = msg.content[0].text.strip()
    return safe_parse_json_v2(raw, context="ats_optimize")


def build_cv_docx(cv_document: dict) -> bytes:
    """
    Builds a real, ATS-readable .docx CV from the structured cv_document.
    No tables, no graphics — plain headings and bullet lists only, so ATS
    parsers can read it cleanly. python-docx is already a real dependency
    of this file (used by _try_docx above for reading); this reuses it for
    writing.
    """
    from docx import Document as DocxWriter
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io as _io

    doc = DocxWriter()
    header = cv_document.get("header", {}) or {}

    name_p = doc.add_paragraph()
    name_run = name_p.add_run(header.get("full_name") or "[Not provided]")
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(header.get("professional_title") or "")
    title_run.font.size = Pt(13)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    contact_bits = [b for b in [header.get("location"), header.get("email"), header.get("phone")]
                    if b and b != "[Not provided]"]
    if contact_bits:
        contact_p = doc.add_paragraph()
        contact_p.add_run(" | ".join(contact_bits)).font.size = Pt(10)
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_heading(text):
        return doc.add_heading(text, level=2)

    if cv_document.get("professional_summary"):
        add_heading("Professional Summary")
        doc.add_paragraph(cv_document["professional_summary"])

    comps = cv_document.get("core_competencies") or []
    if comps:
        add_heading("Core Competencies")
        for i in range(0, len(comps), 2):
            doc.add_paragraph(" • ".join(comps[i:i+2]))

    experience = cv_document.get("professional_experience") or []
    if experience:
        add_heading("Professional Experience")
        for role in experience:
            role_p = doc.add_paragraph()
            role_run = role_p.add_run(f"{role.get('title','')} — {role.get('employer','')}")
            role_run.bold = True
            if role.get("dates"):
                dates_p = doc.add_paragraph()
                dates_run = dates_p.add_run(role["dates"])
                dates_run.italic = True
                dates_run.font.size = Pt(10)
            for bullet in role.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    education = cv_document.get("education") or []
    if education:
        add_heading("Education")
        for edu in education:
            line = edu.get("qualification", "")
            if edu.get("institution"):
                line += f" — {edu['institution']}"
            if edu.get("dates"):
                line += f" ({edu['dates']})"
            doc.add_paragraph(line, style="List Bullet")

    certs = cv_document.get("certifications") or []
    if certs:
        add_heading("Certifications")
        for c in certs:
            doc.add_paragraph(c, style="List Bullet")

    references = cv_document.get("references") or []
    if references:
        add_heading("References")
        for ref in references:
            line = ref.get("name", "")
            bits = [ref.get("title"), ref.get("organisation"), ref.get("contact")]
            bits = [b for b in bits if b]
            if bits:
                line += " — " + ", ".join(bits)
            doc.add_paragraph(line, style="List Bullet")
    elif cv_document.get("references_available_on_request"):
        add_heading("References")
        doc.add_paragraph("Available upon request.")

    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


SCORE_LABELS = {
    "ats_score": "ATS Compatibility",
    "job_match_score": "Job Match",
    "formatting_quality_score": "Formatting Quality",
    "keyword_optimization_score": "Keyword Optimisation",
    "professional_positioning_score": "Professional Positioning",
    "recruiter_readiness_score": "Recruiter Readiness",
    "vacancy_alignment_score": "Vacancy Alignment",
    "evidence_strength_score": "Evidence Strength",
}


def build_ats_report_pdf(meta: dict) -> bytes:
    """
    Generates the downloadable 'ATS Report' — a real PDF built with
    reportlab (pure Python, no LibreOffice/soffice dependency, safe to
    run on Render). Summarises scores, gaps, risks, and the Recruiter's
    First Impression narrative.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    import io as _io

    GOLD = HexColor("#a88820")
    INK  = HexColor("#111417")
    GREY = HexColor("#5b6169")

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=22*mm, bottomMargin=18*mm,
                             leftMargin=20*mm, rightMargin=20*mm)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=INK, fontSize=18, spaceAfter=4)
    eyebrow = ParagraphStyle("eyebrow", parent=styles["Normal"], textColor=GOLD, fontSize=9,
                             spaceAfter=14, alignment=TA_LEFT)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=INK, fontSize=12, spaceBefore=16, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["Normal"], textColor=INK, fontSize=10, leading=15)
    dim = ParagraphStyle("dim", parent=styles["Normal"], textColor=GREY, fontSize=9, leading=13)

    story = []
    story.append(Paragraph("DIRECT LABOUR CONSULT — CAREER INTELLIGENCE REPORT", eyebrow))
    story.append(Paragraph("ATS &amp; Recruiter Readiness Report", h1))
    story.append(Spacer(1, 6))

    hr_score = meta.get("hiring_readiness_score", 0)
    story.append(Paragraph(f"Overall Hiring Readiness Score: <b>{hr_score}/100</b>", body))
    story.append(Spacer(1, 10))

    rows = [["Metric", "Score"]]
    for key, label in SCORE_LABELS.items():
        val = meta.get(key)
        if val is not None:
            rows.append([label, f"{val}/100"])
    tbl = Table(rows, colWidths=[110*mm, 40*mm])
    tbl.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (-1, 0), INK),
        ("TEXTCOLOR", (0, 1), (-1, -1), GREY),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("LINEBELOW", (0, 0), (-1, 0), 0.75, GOLD),
        ("LINEBELOW", (0, 1), (-1, -2), 0.4, HexColor("#e5e7eb")),
    ]))
    story.append(tbl)

    if meta.get("recruiters_first_impression"):
        story.append(Paragraph("Recruiter's First Impression", h2))
        story.append(Paragraph(meta["recruiters_first_impression"], body))

    gaps = meta.get("gaps") or []
    if gaps:
        story.append(Paragraph("Gaps Identified", h2))
        for g in gaps:
            story.append(Paragraph(f"&bull; {g}", body))

    risks = meta.get("risks") or []
    if risks:
        story.append(Paragraph("Risk Flags", h2))
        for r in risks:
            story.append(Paragraph(f"&bull; {r}", body))

    story.append(Spacer(1, 20))
    story.append(Paragraph(
        "Prepared by Direct Labour Consult — HR Advisory, Career Advisory &amp; Recruitment Intelligence.",
        dim
    ))

    doc.build(story)
    return buf.getvalue()


def build_cv_pdf(cv_document: dict) -> bytes:
    """
    Generates the 'Recruiter-Ready CV' as a real PDF (reportlab), mirroring
    build_cv_docx()'s structure so both formats present identically.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    import io as _io

    INK  = HexColor("#111417")
    GOLD = HexColor("#a88820")
    GREY = HexColor("#5b6169")

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                             topMargin=20*mm, bottomMargin=18*mm,
                             leftMargin=20*mm, rightMargin=20*mm)
    styles = getSampleStyleSheet()
    name_style = ParagraphStyle("name", parent=styles["Title"], textColor=INK, fontSize=20,
                                 alignment=TA_CENTER, spaceAfter=2)
    title_style = ParagraphStyle("title", parent=styles["Normal"], textColor=GREY, fontSize=12,
                                  alignment=TA_CENTER, spaceAfter=2)
    contact_style = ParagraphStyle("contact", parent=styles["Normal"], textColor=GREY, fontSize=9,
                                    alignment=TA_CENTER, spaceAfter=14)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=GOLD, fontSize=11, spaceBefore=12, spaceAfter=6)
    body = ParagraphStyle("body", parent=styles["Normal"], textColor=INK, fontSize=10, leading=15)
    bullet = ParagraphStyle("bullet", parent=body, leftIndent=12, bulletIndent=0, spaceAfter=3)
    role_title = ParagraphStyle("role_title", parent=styles["Normal"], textColor=INK, fontSize=10.5,
                                 fontName="Helvetica-Bold", spaceBefore=8, spaceAfter=1)
    role_dates = ParagraphStyle("role_dates", parent=styles["Normal"], textColor=GREY, fontSize=9,
                                 fontName="Helvetica-Oblique", spaceAfter=4)

    header = cv_document.get("header", {}) or {}
    story = [
        Paragraph(header.get("full_name") or "[Not provided]", name_style),
        Paragraph(header.get("professional_title") or "", title_style),
    ]
    contact_bits = [b for b in [header.get("location"), header.get("email"), header.get("phone")]
                    if b and b != "[Not provided]"]
    if contact_bits:
        story.append(Paragraph(" &middot; ".join(contact_bits), contact_style))

    if cv_document.get("professional_summary"):
        story.append(Paragraph("Professional Summary", h2))
        story.append(Paragraph(cv_document["professional_summary"], body))

    comps = cv_document.get("core_competencies") or []
    if comps:
        story.append(Paragraph("Core Competencies", h2))
        story.append(Paragraph(" &nbsp;&bull;&nbsp; ".join(comps), body))

    experience = cv_document.get("professional_experience") or []
    if experience:
        story.append(Paragraph("Professional Experience", h2))
        for role in experience:
            story.append(Paragraph(f"{role.get('title','')} — {role.get('employer','')}", role_title))
            if role.get("dates"):
                story.append(Paragraph(role["dates"], role_dates))
            for b in role.get("bullets", []):
                story.append(Paragraph(f"&bull; {b}", bullet))

    education = cv_document.get("education") or []
    if education:
        story.append(Paragraph("Education", h2))
        for edu in education:
            line = edu.get("qualification", "")
            if edu.get("institution"):
                line += f" — {edu['institution']}"
            if edu.get("dates"):
                line += f" ({edu['dates']})"
            story.append(Paragraph(f"&bull; {line}", bullet))

    certs = cv_document.get("certifications") or []
    if certs:
        story.append(Paragraph("Certifications", h2))
        for c in certs:
            story.append(Paragraph(f"&bull; {c}", bullet))

    references = cv_document.get("references") or []
    if references:
        story.append(Paragraph("References", h2))
        for ref in references:
            line = ref.get("name", "")
            bits = [ref.get("title"), ref.get("organisation"), ref.get("contact")]
            bits = [b for b in bits if b]
            if bits:
                line += " — " + ", ".join(bits)
            story.append(Paragraph(f"&bull; {line}", bullet))
    elif cv_document.get("references_available_on_request"):
        story.append(Paragraph("References", h2))
        story.append(Paragraph("Available upon request.", body))

    doc.build(story)
    return buf.getvalue()


SUPPORTING_DOC_MAX_COUNT = 5
SUPPORTING_DOC_MAX_SIZE  = 10 * 1024 * 1024
SUPPORTING_DOC_ALLOWED   = (".pdf", ".doc", ".docx", ".jpg", ".jpeg", ".png")


def process_supporting_docs(files):
    """
    Optional supporting documents (certifications, qualifications).
    Uses V15B's real extract_cv_text() for PDF/DOC/DOCX (which raises
    ValueError on failure, matching V15B's actual convention) and passes
    images directly to Claude's vision capability. Never fails the main
    request — unsupported/failed files are silently skipped.
    """
    supporting_text_parts = []
    image_blocks = []
    used_count = 0

    for f in files[:SUPPORTING_DOC_MAX_COUNT]:
        if not f or not f.filename:
            continue
        fname = f.filename.lower()
        if not fname.endswith(SUPPORTING_DOC_ALLOWED):
            continue

        f.seek(0, 2)
        fsize = f.tell()
        f.seek(0)
        if fsize > SUPPORTING_DOC_MAX_SIZE:
            continue

        if fname.endswith((".pdf", ".doc", ".docx")):
            try:
                data = f.read()
                text, _method = extract_cv_text(f.filename, data)
                if len(text.strip()) >= 20:
                    supporting_text_parts.append(f"--- {f.filename} ---\n{text.strip()}")
                    used_count += 1
            except ValueError as e:
                log.warning(f"[supporting_docs] extract failed for {f.filename}: {e}")
            except Exception as e:
                log.warning(f"[supporting_docs] unexpected error for {f.filename}: {e}")

        elif fname.endswith((".jpg", ".jpeg", ".png")):
            try:
                raw = f.read()
                media_type = "image/png" if fname.endswith(".png") else "image/jpeg"
                image_blocks.append({
                    "type": "image",
                    "source": {
                        "type": "base64",
                        "media_type": media_type,
                        "data": base64.b64encode(raw).decode("utf-8"),
                    },
                })
                used_count += 1
            except Exception as e:
                log.warning(f"[supporting_docs] image read failed for {f.filename}: {e}")

    supporting_text = "\n\n".join(supporting_text_parts) if supporting_text_parts else ""
    return supporting_text, image_blocks, used_count


@app.route("/generate-pdf-report", methods=["POST"])
def generate_pdf_report():
    """
    Downloads the 'ATS Report (PDF)'. Accepts the meta JSON already
    returned by /optimize-application (cached client-side) — no need to
    re-upload the CV or re-call Claude, since nothing here needs new
    analysis, only a different output format of already-generated data.
    """
    try:
        meta = request.get_json(force=True, silent=True) or {}
        if not meta:
            return jsonify({"success": False, "error": "No report data provided."}), 400
        pdf_bytes = build_ats_report_pdf(meta)
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name="DLC_ATS_Report.pdf",
        )
    except Exception as e:
        log.error(f"[/generate-pdf-report] ERROR: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Could not generate the ATS report. Please try again."}), 500


@app.route("/generate-pdf-cv", methods=["POST"])
def generate_pdf_cv():
    """
    Downloads the 'Recruiter-Ready CV (PDF)'. Accepts the cv_document JSON
    already returned by /optimize-application (cached client-side).
    """
    try:
        body = request.get_json(force=True, silent=True) or {}
        cv_document = body.get("cv_document") or body
        if not cv_document:
            return jsonify({"success": False, "error": "No CV data provided."}), 400
        pdf_bytes = build_cv_pdf(cv_document)
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype="application/pdf",
            as_attachment=True,
            download_name="DLC_Optimised_CV.pdf",
        )
    except Exception as e:
        log.error(f"[/generate-pdf-cv] ERROR: {e}", exc_info=True)
        return jsonify({"success": False, "error": "Could not generate the CV PDF. Please try again."}), 500


@app.route("/optimize-application", methods=["POST"])
def optimize_application():
    """
    DLC ATS CV Generator. Accepts a CV file (required, field name 'file')
    and job description (optional). Generates a full structured CV and a
    real downloadable .docx file, returned via send_file with score/preview
    metadata in the X-DLC-Meta header. Does not touch /analyze,
    /analyze-batch, or any logic above this section.
    """
    ts = datetime.now(timezone.utc).isoformat()
    file = request.files.get("file")
    job_description = (request.form.get("job_description") or "").strip()

    log.info(f"[/optimize-application] admin={check_admin(request)} ts={ts} has_job_desc={bool(job_description)}")

    if not file or not file.filename:
        return jsonify({"success": False, "error": "No CV file uploaded."}), 400

    fname = file.filename.lower()
    if not fname.endswith((".pdf", ".doc", ".docx")):
        return jsonify({"success": False, "error": "Only PDF, DOC, or DOCX files are accepted."}), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_FILE_BYTES:
        return jsonify({"success": False, "error": "File too large. Maximum 10 MB."}), 400

    if job_description and len(job_description) < 40:
        return jsonify({"success": False, "error": "Please provide the full job description, or leave it blank for a general optimisation."}), 400

    try:
        data = file.read()
        try:
            cv_text, method = extract_cv_text(file.filename, data)
        except ValueError as exc:
            return jsonify({"success": False, "error": str(exc)}), 422

        supporting_docs = request.files.getlist("supporting_docs")
        supporting_text, image_blocks, docs_used = process_supporting_docs(supporting_docs)

        user_content = f"CANDIDATE CV:\n{cv_text}"
        if job_description:
            user_content += f"\n\nTARGET JOB DESCRIPTION:\n{job_description}"
        else:
            user_content += "\n\n(No target job description was provided — produce a strong general ATS-optimised CV. Set job_match_score to null and reweight hiring_readiness_score as instructed.)"
        if supporting_text:
            user_content += f"\n\nSUPPORTING DOCUMENTS (certifications/qualifications provided by the candidate):\n{supporting_text}"
        if image_blocks:
            user_content += "\n\n(Additional supporting document images are attached — review them for relevant certifications or qualifications.)"

        result = call_claude_ats(ATS_OPTIMIZE_SYSTEM, user_content, image_blocks=image_blocks or None)

        cv_document = result.get("cv_document", {}) or {}
        docx_bytes = build_cv_docx(cv_document)

        meta = {
            "ats_score":                      result.get("ats_score", 0),
            "job_match_score":                result.get("job_match_score"),
            "formatting_quality_score":        result.get("formatting_quality_score", 0),
            "keyword_optimization_score":      result.get("keyword_optimization_score"),
            "professional_positioning_score":  result.get("professional_positioning_score", 0),
            "recruiter_readiness_score":       result.get("recruiter_readiness_score", 0),
            "vacancy_alignment_score":         result.get("vacancy_alignment_score"),
            "evidence_strength_score":         result.get("evidence_strength_score", 0),
            "hiring_readiness_score":          result.get("hiring_readiness_score", 0),
            "gaps":                            result.get("gaps", []),
            "risks":                           result.get("risks", []),
            "recruiters_first_impression":     result.get("recruiters_first_impression", ""),
            "cv_preview":                      cv_document,
            "keywords_added":                  result.get("keywords_added", []),
            "extraction_method":               method,
            "optimised_at":                    ts,
        }
        if docs_used > 0:
            meta["supporting_docs_note"] = "Additional certifications and qualifications have been integrated"

        log.info(
            f"[/optimize-application] SUCCESS "
            f"hiring_readiness={meta['hiring_readiness_score']} "
            f"job_match={meta['job_match_score']} ats={meta['ats_score']}"
        )

        meta_header = base64.b64encode(json.dumps(meta).encode("utf-8")).decode("ascii")

        response = send_file(
            io.BytesIO(docx_bytes),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="DLC_Optimised_CV.docx",
        )
        response.headers["X-DLC-Meta"] = meta_header
        return response

    except RuntimeError as e:
        log.error(f"[/optimize-application] ERROR: {e}")
        return jsonify({"success": False, "error": str(e)}), 503
    except Exception as e:
        log.error(f"[/optimize-application] ERROR: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "Optimisation could not be completed. Please try again."
        }), 500




# ══════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

# ══════════════════════════════════════════════════════════════════════════
# ── CAREERS APPLICATION — additive feature, isolated from everything above.
#    No database exists in this stack, so submissions are emailed directly
#    to the recruitment inbox with all documents attached. Requires real
#    SMTP credentials set as environment variables (see below) — without
#    them this endpoint will return a clear configuration error rather than
#    silently failing or pretending to have sent something it didn't.
# ══════════════════════════════════════════════════════════════════════════

CAREERS_MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB per document, matches the frontend limit
CAREERS_REQUIRED_FILES = ("cv", "cover_letter_file", "degree_certificate")
CAREERS_OPTIONAL_FILES = ("transcript",)

SMTP_HOST = os.environ.get("SMTP_HOST", "")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")
RECRUITMENT_INBOX = os.environ.get("RECRUITMENT_INBOX", "")


@app.route("/careers/apply", methods=["POST"])
def careers_apply():
    """
    Receives a careers application (text fields + PDF attachments) and
    emails it to the recruitment inbox via SMTP. Requires SMTP_HOST,
    SMTP_PORT, SMTP_USER, SMTP_PASSWORD, and RECRUITMENT_INBOX to be set
    as environment variables on Render — without them, this returns a
    clear 503 configuration error rather than a false "success".
    """
    ts = datetime.now(timezone.utc).isoformat()

    if not (SMTP_HOST and SMTP_USER and SMTP_PASSWORD and RECRUITMENT_INBOX):
        log.error("[/careers/apply] SMTP not configured — missing environment variables")
        return jsonify({
            "success": False,
            "error": "Applications cannot be submitted right now. Please contact Direct Labour Consult directly."
        }), 503

    first_name = (request.form.get("first_name") or "").strip()
    last_name = (request.form.get("last_name") or "").strip()
    email = (request.form.get("email") or "").strip()
    phone = (request.form.get("phone") or "").strip()
    location = (request.form.get("location") or "").strip()
    qualification = (request.form.get("qualification") or "").strip()
    institution = (request.form.get("institution") or "").strip()
    year_graduated = (request.form.get("year_graduated") or "").strip()
    cover_letter_text = (request.form.get("cover_letter_text") or "").strip()
    role = (request.form.get("role") or "Careers Application").strip()

    if not all([first_name, last_name, email, phone, location, qualification, institution, year_graduated]):
        return jsonify({"success": False, "error": "Please complete all required fields."}), 400
    if "@" not in email:
        return jsonify({"success": False, "error": "Please provide a valid email address."}), 400

    attachments = []
    for field in CAREERS_REQUIRED_FILES:
        f = request.files.get(field)
        if not f or not f.filename:
            return jsonify({"success": False, "error": f"Missing required document: {field.replace('_', ' ')}."}), 400
        if not f.filename.lower().endswith(".pdf"):
            return jsonify({"success": False, "error": "All documents must be PDF files."}), 400
        data = f.read()
        if len(data) > CAREERS_MAX_FILE_SIZE:
            return jsonify({"success": False, "error": f"{f.filename} exceeds the 10 MB limit."}), 400
        attachments.append((f.filename, data))

    for field in CAREERS_OPTIONAL_FILES:
        f = request.files.get(field)
        if f and f.filename:
            if not f.filename.lower().endswith(".pdf"):
                return jsonify({"success": False, "error": "All documents must be PDF files."}), 400
            data = f.read()
            if len(data) > CAREERS_MAX_FILE_SIZE:
                return jsonify({"success": False, "error": f"{f.filename} exceeds the 10 MB limit."}), 400
            attachments.append((f.filename, data))

    try:
        msg = EmailMessage()
        msg["Subject"] = f"New Application: {role} — {first_name} {last_name}"
        msg["From"] = SMTP_USER
        msg["To"] = RECRUITMENT_INBOX
        msg["Reply-To"] = email

        body = (
            f"New application received for: {role}\n\n"
            f"Name: {first_name} {last_name}\n"
            f"Email: {email}\n"
            f"Phone: {phone}\n"
            f"Location: {location}\n"
            f"Highest Qualification: {qualification}\n"
            f"Institution: {institution}\n"
            f"Year Graduated: {year_graduated}\n\n"
            f"Cover Letter (text):\n{cover_letter_text or '[Not provided — see attached document]'}\n\n"
            f"Submitted: {ts}\n"
            f"Documents attached: {', '.join(a[0] for a in attachments)}\n"
        )
        msg.set_content(body)

        for filename, data in attachments:
            msg.add_attachment(data, maintype="application", subtype="pdf", filename=filename)

        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)

        log.info(f"[/careers/apply] SUCCESS name={first_name} {last_name} email={email}")
        return jsonify({"success": True})

    except Exception as e:
        log.error(f"[/careers/apply] ERROR: {e}", exc_info=True)
        return jsonify({
            "success": False,
            "error": "We could not submit your application due to a technical issue. Please try again or email us directly."
        }), 500


# ══════════════════════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    log.info("Starting DLC backend %s on port %d", BACKEND_VERSION, port)
    app.run(host="0.0.0.0", port=port, debug=False)
